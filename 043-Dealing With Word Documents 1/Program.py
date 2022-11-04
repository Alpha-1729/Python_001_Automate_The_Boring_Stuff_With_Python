#!/usr/bin/python3
# Dealing with word documents 1

"""
>>>> python-docx documentation
        link: https://python-docx.readthedocs.io
>>>> pip install python-docx
>>>> Download sample docx file.
        link: https://autbor.com/demo.docx
>>>>
>>>>
"""

import docx

d = docx.Document("sample.docx")

# Getting the paragraphs object in the document.
print(d.paragraphs)

# Getting the total number of paragraphs in the documents.
print(len(d.paragraphs))

# Printing the first paragraphs.
print(d.paragraphs[0].text)

# Printing the second paragraphs.
print(d.paragraphs[1].text)

p2 = d.paragraphs[1]

# Each paragraph object has runs, run means the style of the text in the paragraphs.
# Runs occur when the style of the text in the paragraphs changes.
print(p2.runs)

# Run object also has a text member variables.
print(p2.runs[0].text)
print(p2.runs[1].text)  # Bold text in the second paragraph.
print(p2.runs[2].text)
print(p2.runs[3].text)  # Italic text in the second paragraph.

# Check the style of the run in the paragraphs.
print(p2.runs[1].bold)  # Check whether the text is bold or not.
print(p2.runs[3].italic)  # Check whether the text is italic or not.

# Changing the text and style, lets change the text and style of the fourth run.
p2.runs[3].text = "Italic and underline"
p2.runs[3].underline = True

# Each paragraphs has a style like Heading, normal...
# We can also change style of the paragraph.
# p2 is the second paragraph.
# Print the current style of the paragraph 2.
print(p2.style)

# Changing the style of the paragraph.
p2.style = 'Title'

# Saving the edited file.
d.save("edited.docx")

print(" ")
