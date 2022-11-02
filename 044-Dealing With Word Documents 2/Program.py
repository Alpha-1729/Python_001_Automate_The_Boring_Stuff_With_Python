#!/usr/bin/python3
# Dealing with word documents 2

"""
>>>> In this we are creating a new document using the docx.
>>>>
>>>>
>>>>
>>>>
"""

import docx

d = docx.Document()

# Adding two paragraphs.
d.add_paragraph("This is a paragraph")
d.add_paragraph("This is another paragraph")

# Adding a new run in the paragraph 1.
p = d.paragraphs[0]
p.add_run(" This is a new run")

# Printing all the runs in the first paragraphs.
print(p.runs)

# Change the second run in the first paragraph to bold.
p.runs[1].bold = True

d.save("created.docx")

print(" ")
